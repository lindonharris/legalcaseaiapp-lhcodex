# utils/document_loaders/docx_loader.py

import os
import subprocess
import tempfile
from typing import List

from langchain.schema import Document
from .base import BaseDocumentLoader

try:
    import docx  # python-docx
except ImportError:
    raise ImportError("Please install python-docx (pip install python-docx)")

class DocxLoader(BaseDocumentLoader):
    """
    Handles both .doc and .docx files. If given a .doc, it uses
    LibreOffice (headless) to convert it to .docx first, then uses python-docx.
    """

    def load_documents(self, path: str) -> List[Document]:
        """
        path: local filesystem path to a .doc or .docx file.
        Returns a list of langchain.schema.Document, one per non-empty paragraph.
        """
        ext = os.path.splitext(path.lower())[1]
        if ext not in (".doc", ".docx"):
            raise ValueError(f"DocLoader only supports .doc or .docx, but got: {ext}")

        # If it's a .doc, convert to .docx first
        docx_path = path
        temp_docx = None
        if ext == ".doc":
            docx_path = self._convert_doc_to_docx(path)
            temp_docx = docx_path

        # Now docx_path ends with .docx
        try:
            document = docx.Document(docx_path)
        except Exception as e:
            raise RuntimeError(f"python-docx failed to open {docx_path}: {e}")

        results: List[Document] = []
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            results.append(
                Document(
                    page_content=text,
                    metadata={
                        "source_path": os.path.basename(path),
                        "paragraph_index": i
                    }
                )
            )

        # Clean up the temporary .docx if we created one
        if temp_docx and os.path.exists(temp_docx):
            try:
                os.unlink(temp_docx)
            except OSError:
                pass

        return results

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """
        Uses LibreOffice (headless) to convert a .doc → .docx.
        Returns the path to the new .docx file.
        """
        if not os.path.isfile(doc_path):
            raise FileNotFoundError(f"Cannot convert: '{doc_path}' does not exist.")

        # Create a temporary directory for the converted .docx
        tmpdir = tempfile.mkdtemp(prefix="doc_to_docx_")
        convert_cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "docx",
            "--outdir", tmpdir,
            doc_path
        ]

        try:
            subprocess.run(convert_cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"LibreOffice conversion failed for '{doc_path}': {e}")

        base = os.path.splitext(os.path.basename(doc_path))[0]
        converted_path = os.path.join(tmpdir, base + ".docx")
        if not os.path.isfile(converted_path):
            raise RuntimeError(f"Conversion succeeded, but {converted_path} not found.")

        return converted_path
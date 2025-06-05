# utils/document_loaders/docx_loader.py

import os
import subprocess
import tempfile
from typing import List

from langchain.schema import Document
from .base import BaseDocumentLoader

try:
    import docx  # python-docx
except ImportError:
    raise ImportError("Please install python-docx (pip install python-docx)")

class DocxLoader(BaseDocumentLoader):
    """
    Handles both .doc and .docx files. If given a .doc, it uses
    LibreOffice (headless) to convert it to .docx first, then uses python-docx.
    """

    def load_documents(self, path: str) -> List[Document]:
        """
        Returns a list of Document(page_content, metadata) for the given file.
        - If it's a .doc, we try to convert it to .docx (via LibreOffice).
        - If conversion fails, we fallback to extracting plain text into a .txt and return that.
        - If at the end we have .docx, we parse with python-docx.
        - If we have .txt, we read the whole file and split on double-newlines.
        """
        ext = os.path.splitext(path.lower())[1]

        # ─── Step 1: If original was .doc, attempt conversion ───────────────────────
        if ext == ".doc":
            try:
                path = self._convert_doc_to_docx(path)
                ext = os.path.splitext(path.lower())[1]  # now should be ".docx"
            except RuntimeError as e:
                # LibreOffice failed: fall back to a plaintext extractor
                # For example, try antiword (you could also try tika-parser, etc.)
                try:
                    raw = subprocess.check_output(["antiword", path]).decode("utf-8", errors="ignore")
                except Exception as txt_e:
                    raise RuntimeError(f"LibreOffice failed and antiword fallback failed for '{path}': {txt_e}")

                # Write out the raw text to a temporary .txt file
                tmpdir = tempfile.mkdtemp(prefix="doc_to_txt_")
                base = os.path.splitext(os.path.basename(path))[0]
                txt_path = os.path.join(tmpdir, base + ".txt")
                with open(txt_path, "w", encoding="utf-8") as f:
                    f.write(raw)
                path = txt_path
                ext = ".txt"

        # ─── Step 2: NOW handle different extensions ─────────────────────────────────
        ext = os.path.splitext(path.lower())[1]

        # 2A) If we ended up with a .txt from fallback, split on paragraphs:
        if ext == ".txt":
            docs: List[Document] = []
            with open(path, "r", encoding="utf-8") as f:
                entire = f.read()

            # Split on two consecutive newlines (or whatever logic you prefer)
            for i, para in enumerate(entire.split("\n\n")):
                text = para.strip()
                if text:
                    docs.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source_path": os.path.basename(path),
                                "paragraph_index": i,
                            }
                        )
                    )
            return docs

        # 2B) Otherwise, assume it’s now a .docx and parse with python-docx:
        if ext == ".docx":
            from docx import Document as DocxFile

            doc = DocxFile(path)
            docs: List[Document] = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                docs.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source_path": os.path.basename(path),
                            "paragraph_index": i,
                        },
                    )
                )
            return docs

        # 2C) If it’s any other extension, you can either:
        #      – Raise an error
        #      – Or treat .pdf, .txt differently, etc.
        raise ValueError(f"Unsupported file extension: {ext}")

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """
        Uses LibreOffice (headless) to convert a .doc → .docx.
        Returns the path to the new .docx file.
        """
        if not os.path.isfile(doc_path):
            raise FileNotFoundError(f"Cannot convert: '{doc_path}' does not exist.")

        # Create a temporary directory for the converted .docx
        tmpdir = tempfile.mkdtemp(prefix="doc_to_docx_")
        convert_cmd = [
            "libreoffice",
            "--headless",
            "--convert-to", "docx",
            "--outdir", tmpdir,
            doc_path
        ]

        try:
            subprocess.run(convert_cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except subprocess.CalledProcessError:
            # LibreOffice failed: fall back to antiword (or Tika)
            try:
                raw_text = subprocess.check_output(["antiword", doc_path]).decode("utf-8", errors="ignore")
                tmp_txt = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
                tmp_txt.write(raw_text.encode("utf-8"))
                tmp_txt.close()
                return tmp_txt.name
            except Exception as fallback_e:
                raise RuntimeError(f"Both LibreOffice and antiword failed for '{doc_path}': {fallback_e}")

        base = os.path.splitext(os.path.basename(doc_path))[0]
        converted_path = os.path.join(tmpdir, base + ".docx")
        if not os.path.isfile(converted_path):
            raise RuntimeError(f"Conversion succeeded, but {converted_path} not found.")

        return converted_path